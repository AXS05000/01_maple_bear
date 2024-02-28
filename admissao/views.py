import os
import openpyxl
from django.conf import settings
from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic.list import ListView
from django.views.generic.edit import CreateView, UpdateView
from .models import Contrato, Templates, AvaliacaoFDMP
from django.db.models import Q
from .forms import UploadFileForm, AdmissaoForm, FDMPForm, FDMPFormEdit, CNPJForm
from django.http import HttpResponseRedirect
from django.contrib import messages
from django.urls import reverse_lazy
from django.http import FileResponse, HttpResponse, HttpResponseRedirect
from docx2pdf import convert as convert_docx_to_pdf
from docx import Document
from docx2pdf import convert
from subprocess import Popen
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from docx.shared import Pt
from django.views.generic import View

@method_decorator(login_required(login_url="/login/"), name="dispatch")
class ContratoSearchView(ListView):
    model = Contrato
    template_name = "admissao/busca_de_candidatos.html"
    paginate_by = 20

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["q"] = self.request.GET.get("q", "")
        context["order_by"] = self.request.GET.get("order_by", "-id")
        page_obj = context["page_obj"]

        # Obtém o número da página atual
        current_page = page_obj.number

        # Se há mais de 5 páginas
        if page_obj.paginator.num_pages > 5:
            if current_page - 2 < 1:
                start_page = 1
                end_page = 5
            elif current_page + 2 > page_obj.paginator.num_pages:
                start_page = page_obj.paginator.num_pages - 4
                end_page = page_obj.paginator.num_pages
            else:
                start_page = current_page - 2
                end_page = current_page + 2
        else:
            start_page = 1
            end_page = page_obj.paginator.num_pages

        context["page_range"] = range(start_page, end_page + 1)

        return context

    def get_queryset(self):
        query = self.request.GET.get("q")
        order_by = self.request.GET.get("order_by", "-id")
        if query:
            return Contrato.objects.filter(Q(cpf__icontains=query)).order_by(order_by)
        return Contrato.objects.all().order_by(order_by)


@login_required(login_url="/login/")
def upload_template(request):
    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            new_template = Templates(
                name=form.cleaned_data["name"],
                file=request.FILES["file"],
                ano_vigencia=form.cleaned_data["ano_vigencia"],
            )
            new_template.save()
            return HttpResponseRedirect(
                "/upload_template"
            )  # Redirect to a page showing success.
    else:
        form = UploadFileForm()
    return render(request, "admissao/upload.html", {"form": form})


# class FormCandidatoCreateView(CreateView):
#     model = Contrato
#     form_class = AdmissaoForm
#     template_name = "admissao/formulario_candidato.html"
#     success_url = reverse_lazy("form_candidato")

#     def form_valid(self, form):
#         form.instance.created_by = self.request.user
#         return super().form_valid(form)

    # def form_invalid(self, form):
    #     for field, errors in form.errors.items():
    #         for error in errors:
    #             messages.error(
    #                 self.request, f"Erro no campo '{form.fields[field].label}': {error}"
    #             )
    #     return super().form_invalid(form)



class FormCandidatoCreateView(CreateView):
    model = Contrato
    form_class = AdmissaoForm
    template_name = "admissao/formulario_candidato.html"  # substitua com o seu template
    success_url = reverse_lazy(
        "form_candidato"
    )  # substitua com a URL que você quer redirecionar após o sucesso

    def form_valid(self, form):
        cpf = form.cleaned_data.get("cpf")
        collaborator = Contrato.objects.filter(cpf=cpf).first()

        if collaborator:
            # Atualizar o objeto existente
            for field, value in form.cleaned_data.items():
                if (
                    value is not None
                    and hasattr(collaborator, field)
                    and field != "created_by"
                ):
                    setattr(collaborator, field, value)
            collaborator.save()
            self.object = collaborator
        else:
            # Criar um novo objeto
            if self.request.user.is_authenticated:
                form.instance.created_by = self.request.user
            self.object = form.save()
        return HttpResponseRedirect(self.get_success_url())
    
    def form_invalid(self, form):
        for field, errors in form.errors.items():
            for error in errors:
                messages.error(
                    self.request, f"Erro no campo '{form.fields[field].label}': {error}"
                )
        return super().form_invalid(form)



@method_decorator(login_required(login_url="/login/"), name="dispatch")
class CandidatoUpdateView(UpdateView):
    model = Contrato
    form_class = AdmissaoForm
    template_name = "admissao/formulario_candidato_edit.html"
    success_url = reverse_lazy("busca_candidato")


    def form_invalid(self, form):
        for field, errors in form.errors.items():
            for error in errors:
                messages.error(
                    self.request, f"Erro no campo '{form.fields[field].label}': {error}"
                )
        return super().form_invalid(form)

####################GERAÇÃO DO CONTRATO EM PDF##########################


def set_run_text(run, text):
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def generate_contract(template, contrato):
    doc = Document(template.file.path)
    replacements = {k: str(v) for k, v in contrato.get_field_values().items()}

    # Substituição nos parágrafos
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
        if full_text != paragraph.text:  # Se houver substituições
            paragraph.clear()
            set_run_text(paragraph.add_run(), full_text)

    # Substituição nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    for key, value in replacements.items():
                        if key in full_text:
                            full_text = full_text.replace(key, value)
                    if full_text != paragraph.text:  # Se houver substituições
                        paragraph.clear()
                        set_run_text(paragraph.add_run(), full_text)

    # Make sure the contracts directory exists
    contract_directory = os.path.join(settings.MEDIA_ROOT, "contracts")
    os.makedirs(contract_directory, exist_ok=True)

    # Save the new Word document
    new_contract_filename = os.path.join(
        contract_directory, f"{contrato.nome}_{template.name}.docx"
    )
    doc.save(new_contract_filename)

    # Convert the Word document to PDF
    new_contract_pdf_filename = os.path.join(
        contract_directory, f"{contrato.nome}_{template.name}.pdf"
    )
    p = Popen(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            new_contract_filename,
            "--outdir",
            contract_directory,
        ]
    )
    # print("Waiting for conversion...")
    p.wait()
    # print("Conversion finished.")

    # Delete the Word document
    os.remove(new_contract_filename)

    return new_contract_pdf_filename


# def generate_contract(template, contrato):
#     doc = Document(template.file.path)

#     replacements = {k: str(v) for k, v in contrato.get_field_values().items()}

#     # Substituição em parágrafos
#     for paragraph in doc.paragraphs:
#         for key, value in replacements.items():
#             if key in paragraph.text:
#                 paragraph.text = paragraph.text.replace(key, value)

#     # Substituição em tabelas
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for key, value in replacements.items():
#                     if key in cell.text:
#                         cell.text = cell.text.replace(key, value)

#     contract_directory = os.path.join(settings.MEDIA_ROOT, "contracts")
#     os.makedirs(contract_directory, exist_ok=True)

#     new_contract_filename = os.path.join(
#         contract_directory, f"{contrato.nome}_{template.name}.docx"
#     )
#     doc.save(new_contract_filename)

#     return new_contract_filename


def convert_to_pdf(input_filepath, output_filepath):
    convert(input_filepath, output_filepath)


@login_required(login_url="/login/")
def select_contract_id(request, pk):  # Altere 'contrato_id' para 'pk'
    if request.method == "POST":
        template_id = request.POST.get("template")

        contrato = Contrato.objects.get(pk=pk)  # Use 'pk' aqui também
        template = Templates.objects.get(id=template_id)

        contract_filename = generate_contract(template, contrato)

        # Vamos obter apenas o nome do arquivo sem o caminho
        filename = os.path.basename(contract_filename)

        # Retornamos o arquivo como anexo, definindo o nome do arquivo no cabeçalho 'Content-Disposition'
        return FileResponse(
            open(contract_filename, "rb"),
            as_attachment=True,
            filename=filename,
            content_type="application/pdf",
        )

    # Agora, se o request não for POST, o colaborador será buscado pelo parâmetro na URL:
    else:
        contrato = Contrato.objects.get(pk=pk)
        templates = Templates.objects.all()

        return render(
            request,
            "admissao/select_contract_id.html",
            {"contrato": contrato, "templates": templates},
        )
















################################ AVALIAÇÃO FDMP ##################################################
    
@method_decorator(login_required(login_url="/login/"), name="dispatch")
class FormFDMPFormCreateView(CreateView):
    model = AvaliacaoFDMP
    form_class = FDMPForm
    template_name = "admissao/formulario_fdmp.html"
    success_url = reverse_lazy("form_fdmp")

    def form_valid(self, form):
        cpf = form.cleaned_data.get("cpf")
        avaliacao_fdmp = AvaliacaoFDMP.objects.filter(cpf=cpf).first()

        if avaliacao_fdmp:
            # Atualizar o objeto existente
            for field, value in form.cleaned_data.items():
                if value is not None and hasattr(avaliacao_fdmp, field):
                    setattr(avaliacao_fdmp, field, value)
            avaliacao_fdmp.save()
            self.object = avaliacao_fdmp
        else:
            # Criar um novo objeto
            self.object = form.save()

        return HttpResponseRedirect(self.get_success_url())

    def form_invalid(self, form):
        for field, errors in form.errors.items():
            for error in errors:
                messages.error(
                    self.request, f"Erro no campo '{form.fields[field].label}': {error}"
                )
        return super().form_invalid(form)




class FormFDMPUpdateView(UpdateView):
    model = AvaliacaoFDMP
    form_class = FDMPFormEdit
    template_name = "admissao/formulario_fdmp_edit.html"

    def form_valid(self, form):
        self.object = form.save()
        if 'cnpj' in form.changed_data:
            self.object.cnpj = AvaliacaoFDMP.objects.get(pk=self.object.pk).cnpj
            self.object.save()
        # Retorna uma resposta HTTP simples com uma mensagem de sucesso
        return HttpResponse("Formulário enviado com sucesso!", status=200)

    def form_invalid(self, form):
        for field, errors in form.errors.items():
            for error in errors:
                messages.error(
                    self.request, f"Erro no campo '{form.fields[field].label}': {error}"
                )
        return super().form_invalid(form)




@method_decorator(login_required(login_url="/login/"), name="dispatch")
class EscolasSearchView(ListView):
    model = AvaliacaoFDMP
    template_name = "admissao/busca_escolas.html"
    paginate_by = 20

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["q"] = self.request.GET.get("q", "")
        context["order_by"] = self.request.GET.get("order_by", "-id")
        page_obj = context["page_obj"]

        # Obtém o número da página atual
        current_page = page_obj.number

        # Se há mais de 5 páginas
        if page_obj.paginator.num_pages > 5:
            if current_page - 2 < 1:
                start_page = 1
                end_page = 5
            elif current_page + 2 > page_obj.paginator.num_pages:
                start_page = page_obj.paginator.num_pages - 4
                end_page = page_obj.paginator.num_pages
            else:
                start_page = current_page - 2
                end_page = current_page + 2
        else:
            start_page = 1
            end_page = page_obj.paginator.num_pages

        context["page_range"] = range(start_page, end_page + 1)

        return context

    def get_queryset(self):
        query = self.request.GET.get("q")
        order_by = self.request.GET.get("order_by", "-id")
        if query:
            return AvaliacaoFDMP.objects.filter(Q(cnpj__icontains=query)).order_by(order_by)
        return AvaliacaoFDMP.objects.all().order_by(order_by)
    

class CNPJSearchView(View):
    def get(self, request):
        form = CNPJForm()
        return render(request, 'cnpj_search.html', {'form': form})

    def post(self, request):
        form = CNPJForm(request.POST)
        if form.is_valid():
            cnpj = form.cleaned_data['cnpj']
            avaliacao = get_object_or_404(AvaliacaoFDMP, cnpj=cnpj)
            return redirect('form_fdmp_edit', pk=avaliacao.pk)
        return render(request, 'cnpj_search.html', {'form': form})


#################################### IMPORTAÇÃO EXCEL ###############################################


class ExcelImportView(View):
    def get(self, request):
        # Retorna o template de upload de arquivo
        return render(request, 'admissao/upload_excel.html')

    def post(self, request):
        excel_file = request.FILES['excel_file']
        
        # Carrega o arquivo Excel na memória
        workbook = openpyxl.load_workbook(excel_file)
        sheet = workbook.active

        # Itera sobre as linhas do arquivo Excel
        for row in sheet.iter_rows(min_row=2, values_only=True):  # Ignora o cabeçalho
            cnpj = row[0]
            if cnpj:  # Verifica se o CNPJ não está vazio
                # Cria um novo objeto AvaliacaoFDMP
                AvaliacaoFDMP.objects.create(cnpj=cnpj)
        
        return HttpResponse("Importação realizada com sucesso!")
    

